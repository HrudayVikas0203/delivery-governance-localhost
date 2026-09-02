import json
import io
from pathlib import Path
from xml.sax.saxutils import escape

from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile, status
from fastapi.encoders import jsonable_encoder
from sqlalchemy import func, select
from sqlalchemy.orm import Session

from app.core.config import get_settings, resolve_app_path
from app.core.security import get_current_user
from app.db.session import get_db
from app.models.brd import BRDDesignArtifact, BRDDocument, BRDDocumentStatus, BRDRequirementSet
from app.models.delivery import Account, Project
from app.models.people import Employee
from app.schemas.common import (
    BRDArtifactCreate,
    BRDArtifactOut,
    BRDDocumentOut,
    BRDGenerateRequest,
    RequirementOut,
    RequirementSave,
)
from app.services.audit import audit
from app.services.access import require_project_access, require_project_manager, visible_project_ids
from app.services.llm import generate_text
from starlette.responses import Response

router = APIRouter(prefix="/brd", tags=["brd-studio"])
SUPPORTED_ARTIFACT_TYPES = {"business_flow", "architecture"}


def _json_dump(value: object) -> str:
    return json.dumps(value or [], ensure_ascii=False)


def _json_load(value: str | None, fallback: object) -> object:
    if not value:
        return fallback
    try:
        return json.loads(value)
    except json.JSONDecodeError:
        return fallback


def _gemini_json(prompt: str) -> tuple[dict, str]:
    try:
        text, model = generate_text("gemini", prompt)
        cleaned = text.strip()
        if cleaned.startswith("```"):
            cleaned = cleaned.split("\n", 1)[-1].rsplit("```", 1)[0]
        payload = json.loads(cleaned)
        if not isinstance(payload, dict):
            raise ValueError("Gemini response is not an object")
        return payload, model
    except HTTPException:
        raise
    except (ValueError, json.JSONDecodeError) as exc:
        raise HTTPException(status_code=status.HTTP_502_BAD_GATEWAY, detail="Gemini returned an invalid structured artifact.") from exc


def _validate_diagram_payload(artifact_type: str, payload: dict) -> dict:
    if artifact_type == "business_flow":
        nodes = payload.get("nodes")
        edges = payload.get("edges")
        if not isinstance(nodes, list) or len(nodes) < 3 or not isinstance(edges, list):
            raise HTTPException(status_code=502, detail="Gemini returned an incomplete business flow.")
        node_ids = {str(node.get("id")) for node in nodes if isinstance(node, dict) and node.get("id")}
        if len(node_ids) != len(nodes) or any(edge.get("source") not in node_ids or edge.get("target") not in node_ids for edge in edges if isinstance(edge, dict)):
            raise HTTPException(status_code=502, detail="Gemini returned invalid business-flow connections.")
    elif artifact_type == "architecture":
        layers = payload.get("layers")
        if not isinstance(layers, list) or len(layers) < 3:
            raise HTTPException(status_code=502, detail="Gemini returned an incomplete solution architecture.")
        for layer in layers:
            if not isinstance(layer, dict) or not layer.get("name") or not isinstance(layer.get("components"), list):
                raise HTTPException(status_code=502, detail="Gemini returned an invalid solution architecture.")
    return payload


def _artifact_sections(artifact: BRDDesignArtifact) -> list[tuple[str, list[str]]]:
    payload = _json_load(artifact.payload_json, {})
    if not isinstance(payload, dict):
        return [(artifact.title, [str(payload)])]
    sections: list[tuple[str, list[str]]] = []
    if artifact.artifact_type == "business_flow":
        for node in payload.get("nodes", []):
            if isinstance(node, dict):
                detail = node.get("description") or node.get("type") or "Process step"
                sections.append((str(node.get("label") or node.get("id") or "Step"), [str(detail)]))
    else:
        for layer in payload.get("layers", []):
            if isinstance(layer, dict):
                components = [str(item) for item in layer.get("components", [])]
                if layer.get("purpose"):
                    components.insert(0, str(layer["purpose"]))
                sections.append((str(layer.get("name") or "Layer"), components))
    return sections


def _drawio_bytes(artifact: BRDDesignArtifact) -> bytes:
    payload = _json_load(artifact.payload_json, {})
    nodes = []
    edges = []
    if artifact.artifact_type == "business_flow":
        nodes = payload.get("nodes", []) if isinstance(payload, dict) else []
        edges = payload.get("edges", []) if isinstance(payload, dict) else []
    else:
        for index, layer in enumerate(payload.get("layers", []) if isinstance(payload, dict) else []):
            nodes.append({"id": f"layer-{index}", "label": layer.get("name"), "description": ", ".join(layer.get("components", []))})
            if index:
                edges.append({"source": f"layer-{index-1}", "target": f"layer-{index}", "label": "data flow"})
    cells = ['<mxCell id="0"/>', '<mxCell id="1" parent="0"/>']
    ids = set()
    for index, node in enumerate(nodes):
        node_id = str(node.get("id") or f"node-{index}")
        ids.add(node_id)
        value = escape(f"{node.get('label', node_id)}&#xa;{node.get('description', '')}")
        x = 60 + (index % 3) * 280
        y = 60 + (index // 3) * 150
        cells.append(f'<mxCell id="{escape(node_id)}" value="{value}" style="rounded=1;whiteSpace=wrap;html=1;" vertex="1" parent="1"><mxGeometry x="{x}" y="{y}" width="220" height="90" as="geometry"/></mxCell>')
    for index, edge in enumerate(edges):
        source, target = str(edge.get("source", "")), str(edge.get("target", ""))
        if source in ids and target in ids:
            cells.append(f'<mxCell id="edge-{index}" value="{escape(str(edge.get("label", "")))}" style="edgeStyle=orthogonalEdgeStyle;rounded=0;html=1;" edge="1" parent="1" source="{escape(source)}" target="{escape(target)}"><mxGeometry relative="1" as="geometry"/></mxCell>')
    xml = f'<mxfile host="app.diagrams.net"><diagram name="{escape(artifact.title)}"><mxGraphModel><root>{"".join(cells)}</root></mxGraphModel></diagram></mxfile>'
    return xml.encode("utf-8")


def _artifact_export(artifact: BRDDesignArtifact, export_format: str) -> tuple[bytes, str, str]:
    sections = _artifact_sections(artifact)
    safe_stem = "".join(char if char.isalnum() or char in "-_" else "_" for char in artifact.title).strip("_") or "architecture"
    if export_format in {"drawio", "io"}:
        return _drawio_bytes(artifact), "application/vnd.jgraph.mxfile", f"{safe_stem}.drawio"
    if export_format == "pdf":
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

        output = io.BytesIO()
        styles = getSampleStyleSheet()
        document = SimpleDocTemplate(output, pagesize=A4)
        story = [Paragraph(escape(artifact.title), styles["Title"]), Spacer(1, 12)]
        for heading, lines in sections:
            story.extend([Paragraph(escape(heading), styles["Heading2"]), Paragraph(escape(" • ".join(lines)), styles["BodyText"]), Spacer(1, 8)])
        document.build(story)
        return output.getvalue(), "application/pdf", f"{safe_stem}.pdf"
    if export_format == "docx":
        from docx import Document

        document = Document()
        document.add_heading(artifact.title, 0)
        for heading, lines in sections:
            document.add_heading(heading, level=1)
            for line in lines:
                document.add_paragraph(line, style="List Bullet")
        output = io.BytesIO()
        document.save(output)
        return output.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", f"{safe_stem}.docx"
    if export_format == "png":
        from PIL import Image, ImageDraw, ImageFont

        width = 1400
        block_height = 120
        height = max(500, 130 + len(sections) * block_height)
        image = Image.new("RGB", (width, height), "white")
        draw = ImageDraw.Draw(image)
        font = ImageFont.load_default()
        draw.rounded_rectangle((35, 25, width - 35, 90), radius=16, fill="#0f172a")
        draw.text((60, 50), artifact.title, fill="white", font=font)
        for index, (heading, lines) in enumerate(sections):
            y = 115 + index * block_height
            draw.rounded_rectangle((60, y, width - 60, y + 90), radius=12, fill="#eff6ff", outline="#60a5fa", width=2)
            draw.text((85, y + 18), heading, fill="#1d4ed8", font=font)
            draw.text((85, y + 45), " | ".join(lines)[:180], fill="#334155", font=font)
        output = io.BytesIO()
        image.save(output, format="PNG")
        return output.getvalue(), "image/png", f"{safe_stem}.png"
    raise HTTPException(status_code=422, detail="Export format must be pdf, docx, png, or drawio")


def _hydrate_document(document: BRDDocument, db: Session) -> BRDDocument:
    project = db.get(Project, document.project_id)
    setattr(document, "project_name", project.name if project else None)
    return document


def _hydrate_requirement(req: BRDRequirementSet) -> BRDRequirementSet:
    setattr(req, "functional", _json_load(req.functional_json, []))
    setattr(req, "non_functional", _json_load(req.non_functional_json, []))
    setattr(req, "assumptions", _json_load(req.assumptions_json, []))
    return req


def _hydrate_artifact(artifact: BRDDesignArtifact) -> BRDDesignArtifact:
    setattr(artifact, "payload", _json_load(artifact.payload_json, {}))
    return artifact


def _artifact_dict(artifact: BRDDesignArtifact) -> dict:
    hydrated = _hydrate_artifact(artifact)
    return jsonable_encoder({
        "id": hydrated.id,
        "project_id": hydrated.project_id,
        "document_id": hydrated.document_id,
        "artifact_type": hydrated.artifact_type,
        "version": hydrated.version,
        "title": hydrated.title,
        "payload": getattr(hydrated, "payload", {}),
        "ai_provider": hydrated.ai_provider,
        "model_used": hydrated.model_used,
        "created_by_id": hydrated.created_by_id,
        "created_at": hydrated.created_at,
    })


def _requirement_dict(req: BRDRequirementSet) -> dict:
    hydrated = _hydrate_requirement(req)
    return jsonable_encoder({
        "id": hydrated.id,
        "document_id": hydrated.document_id,
        "project_id": hydrated.project_id,
        "version": hydrated.version,
        "overview": hydrated.overview,
        "functional": getattr(hydrated, "functional", []),
        "non_functional": getattr(hydrated, "non_functional", []),
        "assumptions": getattr(hydrated, "assumptions", []),
        "created_by": hydrated.created_by,
        "created_at": hydrated.created_at,
    })


@router.get("/documents", response_model=list[BRDDocumentOut])
def list_documents(
    project_id: str | None = None,
    db: Session = Depends(get_db),
    actor: Employee = Depends(get_current_user),
) -> list[BRDDocument]:
    allowed_ids = visible_project_ids(db, actor)
    if project_id and project_id not in allowed_ids:
        raise HTTPException(status_code=404, detail="Project not found")
    if not allowed_ids:
        return []
    stmt = select(BRDDocument).where(BRDDocument.project_id.in_(allowed_ids)).order_by(BRDDocument.uploaded_at.desc())
    if project_id:
        stmt = stmt.where(BRDDocument.project_id == project_id)
    return [_hydrate_document(document, db) for document in db.scalars(stmt).all()]


@router.post("/documents/upload", response_model=BRDDocumentOut, status_code=201)
async def upload_document(
    project_id: str = Form(...),
    document_type: str = Form("brd"),
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    actor: Employee = Depends(get_current_user),
) -> BRDDocument:
    project = db.get(Project, project_id)
    project = require_project_access(db, actor, project)
    require_project_manager(actor, project, db.get(Account, project.account_id))

    settings = get_settings()
    storage_dir = resolve_app_path(str(settings.report_dir.parent / "brd" / project_id))
    storage_dir.mkdir(parents=True, exist_ok=True)
    content = await file.read(10 * 1024 * 1024 + 1)
    if len(content) > 10 * 1024 * 1024:
        raise HTTPException(status_code=413, detail="BRD file exceeds the 10 MB limit")
    safe_filename = Path(file.filename or "uploaded-brd").name
    document = BRDDocument(
        project_id=project_id,
        filename=safe_filename,
        document_type=document_type,
        content_type=file.content_type,
        size_bytes=len(content),
        status=BRDDocumentStatus.READY,
        uploaded_by_id=actor.id,
        extracted_text=content.decode("utf-8", errors="ignore")[:120000],
    )
    db.add(document)
    db.flush()
    target = storage_dir / f"{document.id}_{document.filename}"
    target.write_bytes(content)
    document.storage_path = str(target)
    audit(db, actor.id, "BRD Uploaded", "BRD Studio", f"BRD {document.filename} uploaded for {project.name}")
    db.commit()
    db.refresh(document)
    return _hydrate_document(document, db)


@router.get("/projects/{project_id}/requirements", response_model=list[RequirementOut])
def list_project_requirements(project_id: str, db: Session = Depends(get_db), actor: Employee = Depends(get_current_user)) -> list[BRDRequirementSet]:
    require_project_access(db, actor, db.get(Project, project_id))
    reqs = db.scalars(
        select(BRDRequirementSet).where(BRDRequirementSet.project_id == project_id).order_by(BRDRequirementSet.created_at.desc())
    ).all()
    return [_hydrate_requirement(req) for req in reqs]


@router.post("/requirements", response_model=RequirementOut, status_code=201)
def save_requirements(
    payload: RequirementSave,
    db: Session = Depends(get_db),
    actor: Employee = Depends(get_current_user),
) -> BRDRequirementSet:
    document = db.get(BRDDocument, payload.document_id)
    if not document:
        raise HTTPException(status_code=404, detail="BRD document not found")
    if document.project_id != payload.project_id:
        raise HTTPException(status_code=422, detail="BRD document does not belong to the selected project")
    project = require_project_access(db, actor, db.get(Project, payload.project_id))
    require_project_manager(actor, project, db.get(Account, project.account_id))
    latest_version = db.scalar(
        select(func.max(BRDRequirementSet.version)).where(BRDRequirementSet.document_id == payload.document_id)
    ) or 0
    req = BRDRequirementSet(
        document_id=payload.document_id,
        project_id=payload.project_id,
        version=latest_version + 1,
        overview=payload.overview,
        functional_json=_json_dump(payload.functional),
        non_functional_json=_json_dump(payload.non_functional),
        assumptions_json=_json_dump(payload.assumptions),
        created_by=payload.created_by,
    )
    db.add(req)
    audit(db, actor.id, "Requirements Saved", "BRD Studio", f"Requirements saved for document {payload.document_id}")
    db.commit()
    db.refresh(req)
    return _hydrate_requirement(req)


@router.get("/projects/{project_id}/artifacts", response_model=list[BRDArtifactOut])
def list_artifacts(
    project_id: str,
    artifact_type: str | None = None,
    db: Session = Depends(get_db),
    actor: Employee = Depends(get_current_user),
) -> list[BRDDesignArtifact]:
    require_project_access(db, actor, db.get(Project, project_id))
    if artifact_type and artifact_type not in SUPPORTED_ARTIFACT_TYPES:
        raise HTTPException(status_code=422, detail="Only business flow and architecture artifacts are supported")
    stmt = select(BRDDesignArtifact).where(
        BRDDesignArtifact.project_id == project_id,
        BRDDesignArtifact.artifact_type.in_(SUPPORTED_ARTIFACT_TYPES),
    ).order_by(BRDDesignArtifact.created_at.desc())
    if artifact_type:
        stmt = stmt.where(BRDDesignArtifact.artifact_type == artifact_type)
    return [_hydrate_artifact(artifact) for artifact in db.scalars(stmt).all()]


@router.get("/artifacts/{artifact_id}", response_model=BRDArtifactOut)
def get_artifact(artifact_id: str, db: Session = Depends(get_db), actor: Employee = Depends(get_current_user)) -> BRDDesignArtifact:
    artifact = db.get(BRDDesignArtifact, artifact_id)
    if not artifact or artifact.artifact_type not in SUPPORTED_ARTIFACT_TYPES:
        raise HTTPException(status_code=404, detail="Artifact not found")
    require_project_access(db, actor, db.get(Project, artifact.project_id))
    return _hydrate_artifact(artifact)


@router.get("/artifacts/{artifact_id}/export")
def export_artifact(
    artifact_id: str,
    format: str,
    db: Session = Depends(get_db),
    actor: Employee = Depends(get_current_user),
) -> Response:
    artifact = db.get(BRDDesignArtifact, artifact_id)
    if not artifact or artifact.artifact_type not in SUPPORTED_ARTIFACT_TYPES:
        raise HTTPException(status_code=404, detail="Artifact not found")
    require_project_access(db, actor, db.get(Project, artifact.project_id))
    content, content_type, filename = _artifact_export(artifact, format.lower())
    return Response(
        content=content,
        media_type=content_type,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.post("/artifacts", response_model=BRDArtifactOut, status_code=201)
def create_artifact(
    payload: BRDArtifactCreate,
    db: Session = Depends(get_db),
    actor: Employee = Depends(get_current_user),
) -> BRDDesignArtifact:
    if payload.artifact_type not in SUPPORTED_ARTIFACT_TYPES:
        raise HTTPException(status_code=422, detail="Only business flow and architecture artifacts are supported")
    project = require_project_access(db, actor, db.get(Project, payload.project_id))
    require_project_manager(actor, project, db.get(Account, project.account_id))
    _validate_diagram_payload(payload.artifact_type, payload.payload)
    if payload.document_id:
        document = db.get(BRDDocument, payload.document_id)
        if not document or document.project_id != payload.project_id:
            raise HTTPException(status_code=422, detail="BRD document does not belong to the selected project")
    latest_version = db.scalar(
        select(func.max(BRDDesignArtifact.version)).where(
            BRDDesignArtifact.project_id == payload.project_id,
            BRDDesignArtifact.artifact_type == payload.artifact_type,
        )
    ) or 0
    artifact = BRDDesignArtifact(
        project_id=payload.project_id,
        document_id=payload.document_id,
        artifact_type=payload.artifact_type,
        version=latest_version + 1,
        title=payload.title,
        payload_json=json.dumps(payload.payload, ensure_ascii=False),
        ai_provider=payload.ai_provider,
        model_used=payload.model_used,
        created_by_id=actor.id,
    )
    db.add(artifact)
    audit(db, actor.id, "BRD Artifact Saved", "BRD Studio", f"{payload.artifact_type} saved for project {payload.project_id}")
    db.commit()
    db.refresh(artifact)
    return _hydrate_artifact(artifact)


@router.get("/documents/{document_id}/artifacts/{artifact_type}/versions", response_model=list[BRDArtifactOut])
def list_artifact_versions(
    document_id: str,
    artifact_type: str,
    db: Session = Depends(get_db),
    actor: Employee = Depends(get_current_user),
) -> list[BRDDesignArtifact]:
    if artifact_type not in SUPPORTED_ARTIFACT_TYPES:
        raise HTTPException(status_code=422, detail="Only business flow and architecture artifacts are supported")
    document = db.get(BRDDocument, document_id)
    if not document:
        raise HTTPException(status_code=404, detail="BRD document not found")
    require_project_access(db, actor, db.get(Project, document.project_id))
    artifacts = db.scalars(
        select(BRDDesignArtifact)
        .where(BRDDesignArtifact.document_id == document_id)
        .where(BRDDesignArtifact.artifact_type == artifact_type)
        .order_by(BRDDesignArtifact.version.desc())
    ).all()
    return [_hydrate_artifact(artifact) for artifact in artifacts]


@router.get("/documents/{document_id}/artifacts/{artifact_type}/compare", response_model=dict)
def compare_artifact_versions(
    document_id: str,
    artifact_type: str,
    v1: int,
    v2: int,
    db: Session = Depends(get_db),
    actor: Employee = Depends(get_current_user),
) -> dict:
    if artifact_type not in SUPPORTED_ARTIFACT_TYPES:
        raise HTTPException(status_code=422, detail="Only business flow and architecture artifacts are supported")
    document = db.get(BRDDocument, document_id)
    if not document:
        raise HTTPException(status_code=404, detail="BRD document not found")
    require_project_access(db, actor, db.get(Project, document.project_id))
    rows = db.scalars(
        select(BRDDesignArtifact)
        .where(BRDDesignArtifact.document_id == document_id)
        .where(BRDDesignArtifact.artifact_type == artifact_type)
        .where(BRDDesignArtifact.version.in_([v1, v2]))
    ).all()
    by_version = {row.version: _json_load(row.payload_json, {}) for row in rows}
    if v1 not in by_version or v2 not in by_version:
        raise HTTPException(status_code=404, detail="One or both versions not found")
    left = json.dumps(by_version[v1], sort_keys=True, indent=2)
    right = json.dumps(by_version[v2], sort_keys=True, indent=2)
    return {
        "documentId": document_id,
        "artifactType": artifact_type,
        "v1": v1,
        "v2": v2,
        "same": left == right,
        "v1Size": len(left),
        "v2Size": len(right),
    }


@router.post("/documents/{document_id}/artifacts/{artifact_type}/restore/{version}", response_model=BRDArtifactOut)
def restore_artifact_version(
    document_id: str,
    artifact_type: str,
    version: int,
    db: Session = Depends(get_db),
    actor: Employee = Depends(get_current_user),
) -> BRDDesignArtifact:
    if artifact_type not in SUPPORTED_ARTIFACT_TYPES:
        raise HTTPException(status_code=422, detail="Only business flow and architecture artifacts are supported")
    target = db.scalar(
        select(BRDDesignArtifact)
        .where(BRDDesignArtifact.document_id == document_id)
        .where(BRDDesignArtifact.artifact_type == artifact_type)
        .where(BRDDesignArtifact.version == version)
    )
    if not target:
        raise HTTPException(status_code=404, detail="Version not found")
    project = require_project_access(db, actor, db.get(Project, target.project_id))
    require_project_manager(actor, project, db.get(Account, project.account_id))
    payload = _json_load(target.payload_json, {})
    restored_payload = BRDArtifactCreate(
        project_id=target.project_id,
        document_id=document_id,
        artifact_type=artifact_type,
        title=f"{target.title} Restored",
        payload=payload if isinstance(payload, dict) else {"payload": payload},
        ai_provider=target.ai_provider,
        model_used=target.model_used,
    )
    return create_artifact(restored_payload, db, actor)


@router.post("/generate", response_model=dict)
def generate_brd_asset(
    payload: BRDGenerateRequest,
    db: Session = Depends(get_db),
    actor: Employee = Depends(get_current_user),
) -> dict:
    project = db.get(Project, payload.project_id)
    project = require_project_access(db, actor, project)
    require_project_manager(actor, project, db.get(Account, project.account_id))
    if payload.artifact_type not in {"requirements", *SUPPORTED_ARTIFACT_TYPES}:
        raise HTTPException(status_code=422, detail="Unsupported BRD artifact type")
    document = db.get(BRDDocument, payload.document_id) if payload.document_id else None
    if document and document.project_id != project.id:
        raise HTTPException(status_code=422, detail="BRD document does not belong to the selected project")
    doc_text = (document.extracted_text if document else "") or payload.prompt or project.description or project.name

    latest_req = db.scalar(
        select(BRDRequirementSet).where(BRDRequirementSet.project_id == project.id).order_by(BRDRequirementSet.version.desc())
    )
    latest_requirements = _hydrate_requirement(latest_req) if latest_req else None
    req_payload = {
        "overview": getattr(latest_requirements, "overview", None) if latest_requirements else None,
        "functional": getattr(latest_requirements, "functional", []) if latest_requirements else [],
        "nonFunctional": getattr(latest_requirements, "non_functional", []) if latest_requirements else [],
    }

    provider_used = "gemini"
    if payload.artifact_type == "requirements":
        result, model_used = _gemini_json(
            "Return only JSON with overview, functional, nonFunctional, assumptions. "
            "Extract requirements from the supplied BRD; do not invent project facts. "
            f"Project: {project.name}\nBRD:\n{doc_text[:12000]}"
        )
        if document:
            saved = save_requirements(
                RequirementSave(
                    document_id=document.id,
                    project_id=project.id,
                    overview=result.get("overview"),
                    functional=result.get("functional", []),
                    non_functional=result.get("nonFunctional", result.get("non_functional", [])),
                    assumptions=result.get("assumptions", []),
                    created_by=provider_used,
                ),
                db,
                actor,
            )
            return {"status": "saved", "provider": provider_used, "model": model_used, "requirements": _requirement_dict(saved)}
        return {"status": "generated", "provider": provider_used, "model": model_used, "requirements": result}

    if payload.artifact_type == "business_flow":
        schema = (
            '{"nodes":[{"id":"stable-id","label":"editable label","description":"specific step",'
            '"type":"actor|process|decision|input|output|exception","actor":"role or system",'
            '"inputs":["..."],"outputs":["..."]}],'
            '"edges":[{"source":"node-id","target":"node-id","label":"condition or outcome",'
            '"kind":"normal|alternate|exception"}]}'
        )
        instructions = (
            "Create a detailed, logically correct and editable business process. Include actors, inputs, outputs, "
            "decisions, system interactions, and applicable alternate/exception paths. Every edge endpoint must "
            "reference a node id. Avoid generic filler."
        )
    else:
        schema = (
            '{"layers":[{"name":"layer","purpose":"specific purpose","components":["component"],'
            '"securityBoundary":"boundary or none"}],'
            '"connections":[{"from":"component","to":"component","label":"protocol/data flow"}],'
            '"deployment":["runtime/infrastructure detail"],"security":["control"],"decisions":["decision"]}'
        )
        instructions = (
            "Create a professional high-level solution architecture that is still technically meaningful. Cover "
            "actors, frontend, API/backend, authentication, business services, data stores, AI/LLM, storage, "
            "external integrations, data flow, security boundaries, and deployment where relevant. Use only "
            "requirements supported by context and do not hard-code this platform's own stack."
        )
    result, model_used = _gemini_json(
        f"Return only valid JSON matching this shape: {schema}\n{instructions}\n"
        f"Project: {project.name}\nDescription: {project.description or ''}\n"
        f"Requirements: {json.dumps(req_payload, ensure_ascii=False)[:12000]}\n"
        f"Additional prompt: {(payload.prompt or '')[:2000]}"
    )
    result = _validate_diagram_payload(payload.artifact_type, result)
    created = create_artifact(
        BRDArtifactCreate(
            project_id=project.id,
            document_id=document.id if document else None,
            artifact_type=payload.artifact_type,
            title=f"{payload.artifact_type.replace('_', ' ').title()} Generated",
            payload=result,
            ai_provider=provider_used,
            model_used=model_used,
        ),
        db,
        actor,
    )
    return {"status": "saved", "provider": provider_used, "model": model_used, "artifact": _artifact_dict(created)}
